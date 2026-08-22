import os
from typing import Optional
import docx
import pdfplumber
from PyPDF2 import PdfReader
from app.core.logging import logger
from app.utils.exceptions import BadRequestException


class TextExtractor:
    """Extracts raw text content from PDF and DOCX resume documents."""

    @classmethod
    def extract_from_pdf(cls, file_path: str) -> str:
        """Extracts text from PDF using pdfplumber with PyPDF2 fallback."""
        text_chunks = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_chunks.append(extracted)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed on {file_path}: {e}. Trying PyPDF2...")
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_chunks.append(extracted)
            except Exception as fallback_error:
                logger.error(f"PyPDF2 fallback also failed on {file_path}: {fallback_error}")
                raise BadRequestException("Failed to extract text from the provided PDF file.")

        full_text = "\n".join(text_chunks).strip()
        if not full_text:
            raise BadRequestException("Uploaded PDF contains no extractable text or is image-only.")
        return full_text

    @classmethod
    def extract_from_docx(cls, file_path: str) -> str:
        """Extracts text paragraphs and table cell contents from DOCX files."""
        try:
            doc = docx.Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))

            result = "\n".join(full_text).strip()
            if not result:
                raise BadRequestException("Uploaded DOCX document contains no readable text.")
            return result
        except Exception as e:
            logger.error(f"DOCX extraction failed on {file_path}: {e}")
            raise BadRequestException(f"Failed to read DOCX file: {str(e)}")

    @classmethod
    def extract(cls, file_path: str, file_type: str) -> str:
        """Routes extraction based on file extension."""
        if not os.path.exists(file_path):
            raise BadRequestException(f"Target file does not exist: {file_path}")

        normalized_type = file_type.lower().replace(".", "").strip()
        if normalized_type == "pdf":
            return cls.extract_from_pdf(file_path)
        elif normalized_type in ["docx", "doc"]:
            return cls.extract_from_docx(file_path)
        else:
            raise BadRequestException(f"Unsupported document format: {file_type}")